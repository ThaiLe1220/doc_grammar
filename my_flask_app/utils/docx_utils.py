""" Filename: docx_utils.py - Directory: my_flask_app/utils
"""
import os
import asyncio
import nltk
from flask import current_app
from docx import Document
from aiohttp import ClientSession
from .grammar_checker import check_grammar
from .reconstructing_sentence import *

# Download necessary Punkt sentence tokenizer if not already downloaded
nltk.download("punkt", quiet=True)


async def correct_text_grammar(file_path):
    """
    Corrects grammar in a DOCX file and applies original formatting to the corrected text.

    :param file_path: The path to the DOCX file to be processed.
    """
    doc = Document(file_path)
    corrections = []  # List to hold correction details
    corrected_paragraphs = (
        {}
    )  # Use a dictionary to map paragraph index to corrected text

    async with ClientSession() as session:
        for i in range(0, len(doc.paragraphs), 10):
            tasks = []
            batch_indices = []  # Store indices of paragraphs in the current batch

            for index, paragraph in enumerate(doc.paragraphs[i : i + 10], start=i):
                if (
                    not paragraph.text.strip()
                    or is_code_snippet(paragraph.text)
                    or paragraph.style.name == "EndNote Bibliography"
                    or paragraph.style.name == "ICCE Affiliations"
                    or paragraph.style.name == "ICCE Author List"
                ):
                    print("Skipping empty or special paragraph.")
                    continue

                task = asyncio.create_task(process_paragraph(paragraph, session))
                tasks.append(task)
                batch_indices.append(
                    index
                )  # Keep track of the paragraph's original index

            # Wait for all tasks in the current batch to complete and store results
            results = await asyncio.gather(*tasks)

            for index, corrected_text in zip(batch_indices, results):
                corrected_paragraphs[
                    index
                ] = corrected_text  # Map index to corrected text

            print("Processed batch of paragraphs. Waiting before next batch...")
            await asyncio.sleep(0.1)

        # Apply the corrected text to each paragraph in the original order
        for index, paragraph in enumerate(doc.paragraphs):
            if (
                paragraph.text.strip()
                and paragraph.style.name != "EndNote Bibliography"
                and index in corrected_paragraphs
            ):
                # paragraph.text = corrected_paragraphs[index]
                correct_paragraph(corrected_paragraphs[index], paragraph)

    # Save the corrected document
    doc.save(file_path)
    print("Completed grammar correction and saved document.")

    # Return all corrections details
    return corrections


async def process_paragraph(paragraph, session):
    corrected_para = ""
    para_word_count = len(paragraph.text.split())

    if para_word_count <= 64:
        # If the paragraph is short enough, process it as a whole
        corrected_para = await async_check_grammar(paragraph.text, session)
        await asyncio.sleep(0.05)  # Delay between batches
        print("Processed short paragraph.")
    else:
        # If the paragraph is long, split into sentences and process in batches of 10
        sentences = nltk.tokenize.sent_tokenize(paragraph.text)
        print(f"Split paragraph into {len(sentences)} sentences.")
        sentence_batches = [sentences[i : i + 10] for i in range(0, len(sentences), 10)]

        for batch in sentence_batches:
            print(f"Processing batch with {len(batch)} sentences.")
            corrected_sentences = await asyncio.gather(
                *[async_check_grammar(sentence, session) for sentence in batch]
            )
            corrected_para += " ".join(corrected_sentences)
            await asyncio.sleep(0.01)  # Delay between batches
            print("Processed a batch of sentences.")

    return corrected_para


def is_code_snippet(paragraph_text):
    return is_java_code_snippet(paragraph_text) or is_html_code_snippet(paragraph_text)


def is_java_code_snippet(paragraph_text):
    # Simple checks for code-like structures
    code_indicators = [
        r"\s*;\s*$",  # lines ending in a semicolon with optional leading whitespace
        r"\s*{\s*$",  # line ending with an opening curly brace with optional leading whitespace
        r"\s*}\s*$",  # line ending with a closing curly brace with optional leading whitespace
        r"\s*class\s+",  # lines that start with 'class' keyword
        r"\s*public\s+",  # lines that start with 'public' keyword
        r"\s*private\s+",  # lines that start with 'private' keyword
        r"\s*(if|for|while|switch|return)\b",  # common control structures
    ]
    for indicator in code_indicators:
        if re.search(indicator, paragraph_text):
            return True
    return False


def is_html_code_snippet(paragraph_text):
    # Simple checks for HTML code-like structures
    html_code_indicators = [
        r"\s*<\s*/?\s*[a-zA-Z][^>]*>",  # Detects HTML tags with optional leading whitespace
        r"\s*<\s*[a-zA-Z]+[^>]*>\s*</\s*[a-zA-Z]+\s*>",  # Detects HTML open and close tag pairs with optional leading whitespace
        r"\s*<!DOCTYPE\s+html>",  # Detects DOCTYPE declaration with optional leading whitespace
        r"\s*<!--.*?-->",  # Detects HTML comments with optional leading whitespace
    ]
    for indicator in html_code_indicators:
        if re.search(indicator, paragraph_text, re.DOTALL | re.MULTILINE):
            return True
    return False


async def async_check_grammar(original_text, session):
    print(f"Checking grammar for text: {original_text[:40]}...")
    corrected_text = await check_grammar(original_text, session)
    print(f"Received corrected text: {corrected_text[:40]}...")  # Print first
    return corrected_text


def correct_paragraph(corrected_para, paragraph):
    print(f"[Corrected Para]: {corrected_para}")
    print(f"[Original Para]: {paragraph}")

    modifications = find_modified_text(paragraph.text, corrected_para)
    paragraph_tokens = custom_tokenize(paragraph.text)
    # print("\nParagraph Tokens:", paragraph_tokens)

    modifications_dict = {item[2]: (item[0], item[1]) for item in modifications}
    # print("\nModifications Dictionary:", modifications_dict)

    current_token_index = 0
    current_pos = 0

    for run_index, run in enumerate(paragraph.runs):
        # print(f"\nProcessing run {run_index} (Text: '{run.text}')")
        new_run_text = ""
        run_start_pos = current_pos
        run_end_pos = current_pos + len(run.text)

        while current_token_index < len(paragraph_tokens) and current_pos < run_end_pos:
            token = paragraph_tokens[current_token_index]
            token_end_pos = current_pos + len(token)
            # print(
            #     f"  Current token: '{token}' (Index: {current_token_index}, Positions: {current_pos}-{token_end_pos})"
            # )

            if current_token_index in modifications_dict:
                original, modified = modifications_dict[current_token_index]
                # print(f"    Modification found: '{original}' -> '{modified}'")

                if token_end_pos <= run_end_pos:
                    # print(f"    Applying full modification to run {run_index}")
                    new_run_text += modified
                    current_token_index += 1
                    current_pos = token_end_pos
                elif run_start_pos < token_end_pos and current_pos < run_end_pos:
                    partial_mod_length = run_end_pos - current_pos
                    # print(
                    #     f"    Applying partial modification to run {run_index}: '{modified[:partial_mod_length]}'"
                    # )
                    new_run_text += modified[:partial_mod_length]

                    if len(original) > partial_mod_length:
                        modifications_dict[current_token_index] = (
                            original[partial_mod_length:],
                            modified[partial_mod_length:],
                        )
                    else:
                        # print("    Skipping token, not in current run")
                        current_token_index += 1

                    current_pos = run_end_pos
                else:
                    break
            else:
                if run_start_pos <= current_pos < run_end_pos:
                    # print(f"    Adding unmodified token: '{token}'")
                    new_run_text += token
                    current_token_index += 1
                    current_pos = token_end_pos
                else:
                    # print("    Token not in current run, moving to next run")
                    break

        run.text = new_run_text
        current_pos = run_end_pos
        # print(f"  Updated run {run_index} text: '{run.text}'")


def find_modified_text(original_text, corrected_text):
    mod = []
    ori_tokens = custom_tokenize(original_text)
    cor_tokens = custom_tokenize(corrected_text)

    ori_matrix = [[token, index] for index, token in enumerate(ori_tokens)]
    cor_matrix = [[token, index] for index, token in enumerate(cor_tokens)]

    temp_mod = find_modified_tokens(ori_matrix, cor_matrix)

    # TODO: replace ori_tokens with mod and then rerun the find_modified_tokens again to improve correctness....
    for item in temp_mod:
        index = item[2]
        ori_tokens[index] = item[1]

    updated_ori_matrix = [[token, index] for index, token in enumerate(ori_tokens)]

    mod = find_modified_tokens(updated_ori_matrix, cor_matrix)
    combined_list = temp_mod + mod

    final_mod = sorted(list(set(combined_list)), key=lambda x: x[2])
    final_mod = [item for item in final_mod if item[0] != item[1]]

    return final_mod


def find_modified_tokens(ori, cor):
    """
    Identifies modifications between original and corrected texts.

    Args:
        ori (list): Original text represented as a list of [token, position] pairs.
        cor (list): Corrected text represented as a list of [token, position] pairs.

    Returns:
        list: A list of tuples of modifications (original_token, modified_token, position).
    """
    ori_index = cor_index = 0
    mod = []

    while ori_index < len(ori) and cor_index < len(cor):
        # print(f"Current Iteration: {ori_index}, {cor_index}")
        ori_token, ori_pos = ori[ori_index]
        cor_token, _ = cor[cor_index]

        if ori_token == cor_token:
            # Tokens are identical, move to the next pair
            ori_index += 1
            cor_index += 1
        else:
            if ori_index + 4 < len(ori) and cor_index + 4 < len(cor):
                if ori[ori_index + 4][0] == cor[cor_index + 2][0]:
                    mod.append((ori_token, "", ori_pos))
                    mod.append((" ", "", ori_pos + 1))
                    mod.append((ori[ori_index + 2][0], cor_token, ori_pos + 2))

                    ori_index += 3
                    cor_index += 1
                    continue
                elif ori[ori_index + 2][0] == cor[cor_index + 2][0]:
                    mod.append((ori_token, cor_token, ori_pos))
                    ori_index += 1
                    cor_index += 1
                    continue
                elif ori[ori_index + 2][0] == cor[cor_index + 4][0]:
                    mod.append(
                        (ori_token, cor_token + " " + cor[cor_index + 2][0], ori_pos)
                    )
                    ori_index += 1
                    cor_index += 3
                    continue
                elif (
                    ori[ori_index + 2][0] != cor[cor_index + 2][0]
                    and ori[ori_index + 4][0] == cor[cor_index + 4][0]
                ):
                    mod.append((ori_token, cor_token, ori_pos))
                    mod.append(
                        (ori[ori_index + 2][0], cor[cor_index + 2][0], ori_pos + 2)
                    )
                    ori_index += 3
                    cor_index += 3

                else:
                    ori_index += 1
                    cor_index += 1
            else:
                # Handle end-of-text cases
                if len(ori) - ori_index == len(cor) - cor_index:
                    mod.append((ori_token, cor_token, ori_pos))
                    ori_index += 1
                    cor_index += 1
                elif len(ori) - ori_index > len(cor) - cor_index:
                    mod.append((ori_token, "", ori_pos))
                    mod.append((" ", "", ori_pos + 1))
                    mod.append((ori[ori_index + 2][0], cor_token, ori_pos + 2))
                    ori_index += 3
                    cor_index += 1
                else:
                    mod.append(
                        (
                            ori_token,
                            cor_token + " " + cor[cor_index + 2][0],
                            ori_pos,
                        )
                    )
                    ori_index += 1
                    cor_index += 3

    return mod
